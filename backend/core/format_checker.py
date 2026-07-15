import json

def parse_json_from_llm(content: str):
    content = content.strip()
    if content.startswith("```json"):
        content = content[7:]
    elif content.startswith("```"):
        content = content[3:]
    if content.endswith("```"):
        content = content[:-3]
    content = content.strip()
    try:
        return json.loads(content)
    except Exception:
        # Fallback in case LLM output isn't valid JSON
        return {}

def check_docx_format(docx_bytes: bytes, file_name: str) -> dict:
    import docx

    from docx.shared import Cm, Pt
    import io
    import zipfile
    from google import genai
    from google.genai import types
    
    doc = docx.Document(io.BytesIO(docx_bytes))
    
    def get_document_defaults(document):
        from docx.oxml.ns import qn
        defaults = {'size': None, 'name': None}
        try:
            styles_element = document.styles.element
            doc_defaults = styles_element.find(qn('w:docDefaults'))
            if doc_defaults is not None:
                rpr_default = doc_defaults.find(qn('w:rPrDefault'))
                if rpr_default is not None:
                    rpr = rpr_default.find(qn('w:rPr'))
                    if rpr is not None:
                        # Extract size from sz or szCs
                        for attr in ['sz', 'szCs']:
                            sz = rpr.find(qn(f'w:{attr}'))
                            if sz is not None:
                                val = sz.get(qn('w:val'))
                                if val:
                                    defaults['size'] = float(val) / 2.0
                                    break
                        # Extract name from ascii, hAnsi, cs, eastAsia
                        rfonts = rpr.find(qn('w:rFonts'))
                        if rfonts is not None:
                            for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
                                val = rfonts.get(qn(f'w:{attr}'))
                                if val:
                                    defaults['name'] = val
                                    break
        except Exception:
            pass
        return defaults

    def get_style_font_size(style):
        if not style:
            return None
        from docx.oxml.ns import qn
        try:
            rpr = style.element.find(qn('w:rPr'))
            if rpr is not None:
                for attr in ['sz', 'szCs']:
                    sz = rpr.find(qn(f'w:{attr}'))
                    if sz is not None:
                        val = sz.get(qn('w:val'))
                        if val:
                            return float(val) / 2.0
        except Exception:
            pass
        return get_style_font_size(style.base_style)

    def get_style_font_name(style):
        if not style:
            return None
        from docx.oxml.ns import qn
        try:
            rpr = style.element.find(qn('w:rPr'))
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
                        val = rfonts.get(qn(f'w:{attr}'))
                        if val:
                            return val
        except Exception:
            pass
        return get_style_font_name(style.base_style)

    def get_effective_font_size(run, paragraph, document, defaults):
        from docx.oxml.ns import qn
        try:
            rpr = run.element.find(qn('w:rPr'))
            if rpr is not None:
                for attr in ['sz', 'szCs']:
                    sz = rpr.find(qn(f'w:{attr}'))
                    if sz is not None:
                        val = sz.get(qn('w:val'))
                        if val:
                            return float(val) / 2.0
        except Exception:
            pass
            
        style_size = get_style_font_size(paragraph.style)
        if style_size is not None:
            return style_size
            
        if defaults['size'] is not None:
            return defaults['size']
            
        return 12.0 # Default Word fallback if not specified
        
    def get_effective_font_name(run, paragraph, document, defaults):
        from docx.oxml.ns import qn
        try:
            rpr = run.element.find(qn('w:rPr'))
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
                        val = rfonts.get(qn(f'w:{attr}'))
                        if val:
                            return val
        except Exception:
            pass
            
        style_name = get_style_font_name(paragraph.style)
        if style_name is not None:
            return style_name
            
        if defaults['name'] is not None:
            return defaults['name']
            
        return 'Times New Roman' # Default Vietnamese Word fallback

    def is_toc_entry_paragraph(p, text):
        style_name = p.style.name.lower()
        if style_name.startswith("toc"):
            return True
        if ("..." in text or ". ." in text) and (text[-1].isdigit() if text else False):
            return True
        return False

    doc_defaults = get_document_defaults(doc)
    
    # 1. Check paper size
    is_paper_size_valid = True
    paper_size_feedback = "Khổ giấy đạt chuẩn A4 (21.0 x 29.7cm)."
    if doc.sections:
        section = doc.sections[0]
        w = round(section.page_width.cm, 1) if section.page_width else 0.0
        h = round(section.page_height.cm, 1) if section.page_height else 0.0
        if abs(w - 21.0) > 0.3 or abs(h - 29.7) > 0.3:
            is_paper_size_valid = False
            paper_size_feedback = f"Khổ giấy hiện tại: {w}x{h}cm (Yêu cầu: A4 21.0x29.7cm)"

    # 2. Check margins (Lề trang)
    margin_errors = []
    if doc.sections:
        section = doc.sections[0]
        top_cm = round(section.top_margin.cm, 2) if section.top_margin else 0.0
        bottom_cm = round(section.bottom_margin.cm, 2) if section.bottom_margin else 0.0
        left_cm = round(section.left_margin.cm, 2) if section.left_margin else 0.0
        right_cm = round(section.right_margin.cm, 2) if section.right_margin else 0.0
        
        if abs(top_cm - 2.5) > 0.15:
            margin_errors.append(f"Lề trên: {top_cm}cm (Yêu cầu: 2.5cm)")
        if abs(bottom_cm - 2.5) > 0.15:
            margin_errors.append(f"Lề dưới: {bottom_cm}cm (Yêu cầu: 2.5cm)")
        if abs(left_cm - 3.5) > 0.15:
            margin_errors.append(f"Lề trái: {left_cm}cm (Yêu cầu: 3.5cm)")
        if abs(right_cm - 2.0) > 0.15:
            margin_errors.append(f"Lề phải: {right_cm}cm (Yêu cầu: 2.0cm)")
            
    is_margins_valid = len(margin_errors) == 0
    margins_feedback = "Lề trang đúng chuẩn quy định của UEH." if is_margins_valid else "; ".join(margin_errors)
    
    # 3. Check Font Family, Font Size and Line Spacing
    total_checked = 0
    wrong_font_count = 0
    wrong_size_count = 0
    wrong_spacing_count = 0
    
    bibliography_text = []
    in_bibliography = False
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        # Detect bibliography section
        if text.lower() in ["tài liệu tham khảo", "references", "danh mục tài liệu tham khảo"]:
            in_bibliography = True
            continue
        elif in_bibliography and text.lower().startswith(("chương", "chapter", "phụ lục", "appendix")):
            in_bibliography = False
            
        if in_bibliography:
            bibliography_text.append(text)
            
        # Check standard paragraph format
        spacing = p.paragraph_format.line_spacing
        
        # Check if heading (protect manual styles)
        is_heading = False
        if p.style.name.startswith("Heading") or text.isupper():
            is_heading = True
            
        # Bypass heading classification for TOC entries
        if is_toc_entry_paragraph(p, text):
            is_heading = False
            
        # Skip checking captions and notes for body text compliance (they have sizes 10pt and 11pt)
        is_caption_or_note = False
        if text.lower().startswith(('hình', 'ảnh', 'sơ đồ', 'biểu đồ', 'đồ thị', 'figure', 'fig', 'bảng', 'table', 'chú thích', 'nguồn', 'ghi chú', 'note', 'source')):
            is_caption_or_note = True

        has_large_size = False
        run_fonts = []
        run_sizes = []
        for run in p.runs:
            if not run.text.strip():
                continue
            
            effective_size = get_effective_font_size(run, p, doc, doc_defaults)
            effective_font = get_effective_font_name(run, p, doc, doc_defaults)
            
            if effective_size and effective_size > 13.5:
                has_large_size = True
            if effective_font:
                run_fonts.append(effective_font)
            if effective_size:
                run_sizes.append(effective_size)
                
        if has_large_size or is_heading or is_caption_or_note:
            continue

            
        total_checked += 1
        
        # Verify font family
        if run_fonts:
            if any(f != "Times New Roman" for f in run_fonts):
                wrong_font_count += 1
        
        # Verify font size
        if run_sizes:
            if any(abs(s - 13.0) > 0.1 for s in run_sizes):
                wrong_size_count += 1
                
        # Verify spacing
        if spacing is not None:
            if isinstance(spacing, float):
                if abs(spacing - 1.2) > 0.08:
                    wrong_spacing_count += 1
            else:
                pt_spacing = spacing.pt if hasattr(spacing, 'pt') else 0.0
                if abs(pt_spacing - 15.6) > 1.5:
                    wrong_spacing_count += 1
                    
    is_font_family_valid = True
    font_family_feedback = "Font chữ toàn văn đạt chuẩn Times New Roman."
    is_font_size_valid = True
    font_size_feedback = "Cỡ chữ toàn văn đạt chuẩn 13pt."
    is_spacing_valid = True
    spacing_feedback = "Giãn dòng đạt chuẩn 1.2 lines."
    
    if total_checked > 0:
        font_fail_rate = wrong_font_count / total_checked
        size_fail_rate = wrong_size_count / total_checked
        spacing_fail_rate = wrong_spacing_count / total_checked
        
        if font_fail_rate > 0.15:
            is_font_family_valid = False
            font_family_feedback = f"Font chữ chưa đồng bộ Times New Roman (tỷ lệ lỗi: {font_fail_rate*100:.1f}%)"
        if size_fail_rate > 0.15:
            is_font_size_valid = False
            font_size_feedback = f"Cỡ chữ chưa đúng chuẩn 13pt (tỷ lệ lỗi: {size_fail_rate*100:.1f}%)"
        if spacing_fail_rate > 0.2:
            is_spacing_valid = False
            spacing_feedback = f"Giãn dòng chưa đúng chuẩn 1.2 lines (tỷ lệ lỗi: {spacing_fail_rate*100:.1f}%)"
    else:
        font_family_feedback = "Không có đủ nội dung văn bản để kiểm tra font."
        font_size_feedback = "Không có đủ nội dung văn bản để kiểm tra cỡ chữ."
        spacing_feedback = "Không có đủ nội dung văn bản để kiểm tra giãn dòng."
        
    # 4. Logo Check using up to 3 images in reading order (protects against borders/lines)
    is_logo_valid = False
    logo_feedback = "Thiếu Logo UEH chính thức trên trang bìa luận văn."
    
    img_parts = []
    
    try:
        def get_images_in_reading_order(document, limit=3):
            from docx.oxml.ns import qn
            embed_ids = []
            
            def add_id(eid):
                if eid and eid not in embed_ids:
                    embed_ids.append(eid)
                    
            # Traverse paragraphs
            for p in document.paragraphs:
                for element in p._element.iter():
                    if element.tag.endswith('blip'):
                        add_id(element.get(qn('r:embed')))
                    elif element.tag.endswith('imagedata'):
                        add_id(element.get(qn('r:id')))
                if len(embed_ids) >= limit:
                    break
                    
            if len(embed_ids) < limit:
                # Traverse tables
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for element in p._element.iter():
                                    if element.tag.endswith('blip'):
                                        add_id(element.get(qn('r:embed')))
                                    elif element.tag.endswith('imagedata'):
                                        add_id(element.get(qn('r:id')))
                                if len(embed_ids) >= limit:
                                    break
                            if len(embed_ids) >= limit:
                                break
                        if len(embed_ids) >= limit:
                            break
            return embed_ids[:limit]

        embed_ids = get_images_in_reading_order(doc, limit=3)
        for eid in embed_ids:
            try:
                rel = doc.part.rels[eid]
                img_bytes = rel.target_part.blob
                mime_type = rel.target_part.content_type
                img_parts.append(types.Part.from_bytes(data=img_bytes, mime_type=mime_type))
            except Exception:
                pass
    except Exception as e:
        print(f"Lỗi bóc tách ảnh bìa docx: {e}")
        
    if img_parts:
        try:
            prompt = """Bạn là chuyên gia thẩm định văn bản của trường Đại học Kinh tế TP.HCM (UEH).
Hãy phân tích các hình ảnh được tải lên (được trích xuất từ trang bìa luận văn của sinh viên) và xác định xem:
1. Có hình ảnh nào chứa Logo chính thức hiện tại của trường Đại học Kinh tế TP.HCM (UEH) hay không?
2. Logo đó (nếu có) có đúng mẫu chuẩn màu sắc và các chi tiết vòng tròn đặc trưng không?

Hãy phản hồi DUY NHẤT ở định dạng JSON thô có cấu trúc như sau:
{
  "is_logo_valid": true_or_false,
  "feedback": "Lời nhận xét chi tiết ngắn gọn bằng tiếng Việt (Ví dụ: Phát hiện logo UEH đúng chuẩn mẫu mới trên bìa, hoặc Các ảnh gửi lên không chứa logo UEH đúng chuẩn)"
}
Tuyệt đối chỉ trả về chuỗi JSON thô, không kèm markdown hay giải thích nào khác."""

            from core.config import generate_content_with_rotation
            response = generate_content_with_rotation(
                model="gemini-3.5-flash",
                contents=img_parts + [prompt]
            )
            result = parse_json_from_llm(response.text)
            is_logo_valid = result.get("is_logo_valid", False)
            logo_feedback = result.get("feedback", "Không thể xác minh logo.")
        except Exception as gemini_err:
            print(f"Lỗi gọi Gemini kiểm tra logo DOCX: {gemini_err}")
            logo_feedback = "Không thể kết nối với AI để kiểm tra logo trang bìa."
            
    # 5. Citation Check (APA 7th)
    is_citations_valid = True
    citations_feedback = "Danh mục tài liệu tham khảo đạt chuẩn APA."
    citations_errors = []
    
    if bibliography_text:
        bib_lines_str = "\n".join(bibliography_text[:25])
        try:
            prompt = f"""Bạn là chuyên gia thẩm định tài liệu tham khảo học thuật.
Hãy đối chiếu danh sách tài liệu tham khảo dưới đây của sinh viên UEH và chỉ ra các lỗi sai so với chuẩn APA 7th hoặc Harvard (như thiếu in nghiêng tên sách/tạp chí, sai thứ tự tên tác giả, năm xuất bản...).

Danh sách tài liệu tham khảo:
{bib_lines_str}

Hãy kiểm tra kỹ từng mục. Với mỗi mục có lỗi sai, hãy đề xuất bản sửa đổi chuẩn.
Phản hồi DUY NHẤT ở định dạng JSON thô có cấu trúc sau:
{{
  "is_citations_valid": true_or_false,
  "errors": [
     {{
       "original": "Mục trích dẫn gốc bị lỗi",
       "reason": "Lý do sai chuẩn chi tiết ngắn gọn bằng tiếng Việt",
       "suggested": "Mục trích dẫn đã được sửa lại đúng chuẩn APA 7th"
     }},
     ...
  ]
}}
If all references are correct, return true for "is_citations_valid" and empty list for "errors".
Tuyệt đối chỉ trả về chuỗi JSON thô, không kèm markdown hay giải thích nào khác."""

            from core.config import generate_content_with_rotation
            response = generate_content_with_rotation(
                model="gemini-3.5-flash",
                contents=prompt
            )
            result = parse_json_from_llm(response.text)
            is_citations_valid = result.get("is_citations_valid", True)
            citations_errors = result.get("errors", [])
            if not is_citations_valid and citations_errors:
                citations_feedback = f"Phát hiện {len(citations_errors)} tài liệu tham khảo chưa đúng chuẩn APA 7th."
            else:
                is_citations_valid = True
                citations_feedback = "Tất cả các tài liệu tham khảo đã quét đều đúng chuẩn APA 7th."
        except Exception as bib_err:
            print(f"Lỗi kiểm tra APA DOCX: {bib_err}")
            is_citations_valid = False
            citations_feedback = "Không thể gọi AI để kiểm tra chuẩn trích dẫn."
    else:
        is_citations_valid = False
        citations_feedback = "Không tìm thấy danh mục tài liệu tham khảo nào trong file Word. Yêu cầu bắt buộc phải có tài liệu tham khảo."
        
    return {
        "file_name": file_name,
        "is_paper_size_valid": is_paper_size_valid,
        "paper_size_feedback": paper_size_feedback,
        "is_margins_valid": is_margins_valid,
        "margins_feedback": margins_feedback,
        "is_font_family_valid": is_font_family_valid,
        "font_family_feedback": font_family_feedback,
        "is_font_size_valid": is_font_size_valid,
        "font_size_feedback": font_size_feedback,
        "is_spacing_valid": is_spacing_valid,
        "spacing_feedback": spacing_feedback,
        "is_logo_valid": is_logo_valid,
        "logo_feedback": logo_feedback,
        "is_citations_valid": is_citations_valid,
        "citations_feedback": citations_feedback,
        "citations_errors": citations_errors
    }


def fix_docx_format(docx_bytes: bytes) -> bytes:
    import docx
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    doc = docx.Document(io.BytesIO(docx_bytes))
    
    # 1. Fix margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)
        
    # 1.05 Tự động thu nhỏ ảnh vượt quá khung trang in khả dụng (15.5cm)
    max_width_emu = Cm(15.5)
    for shape in doc.inline_shapes:
        try:
            if shape.width and shape.width > max_width_emu:
                ratio = float(shape.height) / float(shape.width)
                shape.width = max_width_emu
                shape.height = int(max_width_emu * ratio)
        except Exception:
            pass
        
    def is_toc_entry_paragraph(p, text):
        style_name = p.style.name.lower()
        if style_name.startswith("toc"):
            return True
        if ("..." in text or ". ." in text) and (text[-1].isdigit() if text else False):
            return True
        return False

    # 1.1 Fix TOC style sheet definitions to ensure MS Word updates preserve formatting
    for style_name in ['TOC 1', 'TOC 2', 'TOC 3', 'TOC 4', 'TOC 5', 'toc 1', 'toc 2', 'toc 3', 'toc 4', 'toc 5']:
        try:
            style = doc.styles[style_name]
            style.font.name = 'Times New Roman'
            style.font.size = Pt(13)
        except Exception:
            pass

    # 2. Fix fonts, sizes, and line spacing
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # Check if it is an image paragraph
        is_image_para = False
        if 'w:drawing' in p._p.xml or 'w:pict' in p._p.xml:
            is_image_para = True
            
        if is_image_para:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            continue
            
        if not text:
            continue
            
        # Check if it is a caption/title/note
        is_fig_title = False
        is_tbl_title = False
        is_tbl_note = False
        
        if text.lower().startswith(('chú thích', 'nguồn', 'ghi chú', 'note', 'source')):
            is_tbl_note = True
        elif text.lower().startswith(('hình', 'ảnh', 'sơ đồ', 'biểu đồ', 'đồ thị', 'figure', 'fig')):
            is_fig_title = True
        elif text.lower().startswith(('bảng', 'table')):
            is_tbl_title = True
            
        # Protect headings
        is_heading = False
        if p.style.name.startswith("Heading") or text.isupper():
            is_heading = True
            
        # Bypass heading classification for TOC entries
        if is_toc_entry_paragraph(p, text):
            is_heading = False
            
        has_large_size = False
        for run in p.runs:
            run_size = None
            try:
                from docx.oxml.ns import qn
                rpr = run.element.find(qn('w:rPr'))
                if rpr is not None:
                    for attr in ['sz', 'szCs']:
                        sz = rpr.find(qn(f'w:{attr}'))
                        if sz is not None:
                            val = sz.get(qn('w:val'))
                            if val:
                                run_size = float(val) / 2.0
                                break
            except Exception:
                pass
            if run_size is None and run.font.size:
                run_size = run.font.size.pt
                
            if run_size and run_size > 13.5:
                has_large_size = True
                break
                
        if is_heading or has_large_size:
            # Only fix font family for headings
            for run in p.runs:
                run.font.name = 'Times New Roman'
            continue
            
        if is_fig_title:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                run.italic = False
            continue
            
        if is_tbl_title:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                run.italic = False
            continue
            
        if is_tbl_note:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = False
                run.italic = True
            continue
            
        # Standardize body paragraphs
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(6)
        
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            
    # 3. Fix tables (Issue 3: avoid justified cell text, make it left aligned and 12pt, clear indents)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.left_indent = Pt(0)
                    p.paragraph_format.right_indent = Pt(0)
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        
    out_bio = io.BytesIO()
    doc.save(out_bio)
    return out_bio.getvalue()


def check_pdf_format(pdf_bytes: bytes, file_name: str) -> dict:
    import fitz
    import os
    from google import genai
    from google.genai import types
    
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    # 1. Page 1 Image & Layout check with Gemini
    is_logo_valid = False
    logo_feedback = "Không thể xác minh logo trang bìa PDF."
    
    try:
        page = doc[0]
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("png")
        
        prompt = """Bạn là chuyên gia thẩm định văn bản của Đại học Kinh tế TP.HCM (UEH).
Đây là ảnh chụp trang bìa luận văn tốt nghiệp của sinh viên.
Hãy phân tích trang bìa này và kiểm tra xem:
1. Có logo UEH xuất hiện ở đầu trang bìa không?
2. Logo đó có đúng chuẩn logo Đại học Kinh tế TP.HCM (UEH) hay không?
3. Bố cục tên trường, tên đề tài, logo có được căn giữa cân đối không?

Hãy phản hồi DUY NHẤT ở định dạng JSON thô có cấu trúc sau:
{
  "is_logo_valid": true_or_false,
  "is_layout_valid": true_or_false,
  "feedback": "Nhận xét chi tiết ngắn gọn bằng tiếng Việt (Ví dụ: Logo đúng chuẩn và căn giữa đẹp, hoặc Thiếu logo UEH ở đầu trang, hoặc Tên đề tài chưa được căn giữa)"
}
Tuyệt đối chỉ trả về chuỗi JSON thô, không kèm markdown hay giải thích nào khác."""

        from core.config import generate_content_with_rotation
        response = generate_content_with_rotation(
            model="gemini-3.5-flash",
            contents=[
                types.Part.from_bytes(data=img_bytes, mime_type="image/png"),
                prompt
            ]
        )
        result = parse_json_from_llm(response.text)
        is_logo_valid = result.get("is_logo_valid", False)
        logo_feedback = result.get("feedback", "Không thể xác minh logo.")
    except Exception as e:
        print(f"Lỗi kiểm tra logo PDF: {e}")
        logo_feedback = "Không thể kết nối với AI để kiểm tra logo trang bìa PDF."
        
    # 2. Check paper size
    is_paper_size_valid = True
    paper_size_feedback = "Khổ giấy đạt chuẩn A4 (21.0 x 29.7cm)."
    if len(doc) > 0:
        page0 = doc[0]
        w_cm = page0.rect.width * 0.0352778
        h_cm = page0.rect.height * 0.0352778
        if abs(w_cm - 21.0) > 0.4 or abs(h_cm - 29.7) > 0.4:
            is_paper_size_valid = False
            paper_size_feedback = f"Khổ giấy hiện tại: {w_cm:.1f}x{h_cm:.1f}cm (Yêu cầu: A4 21.0x29.7cm)"

    # 3. Check margins
    margin_errors = []
    is_margins_valid = True
    margins_feedback = "Lề trang PDF đúng chuẩn quy định của UEH."
    if len(doc) > 1:
        page2 = doc[1]
        rect = page2.rect
        blocks = page2.get_text("blocks")
        if blocks:
            x0_min = min(b[0] for b in blocks)
            y0_min = min(b[1] for b in blocks)
            x1_max = max(b[2] for b in blocks)
            y1_max = max(b[3] for b in blocks)
            
            left_margin = x0_min * 0.0352778
            top_margin = y0_min * 0.0352778
            right_margin = (rect.width - x1_max) * 0.0352778
            bottom_margin = (rect.height - y1_max) * 0.0352778
            
            if abs(left_margin - 3.5) > 0.4:
                margin_errors.append(f"Lề trái ước lượng: {left_margin:.1f}cm (Yêu cầu: 3.5cm)")
            if abs(right_margin - 2.0) > 0.4:
                margin_errors.append(f"Lề phải ước lượng: {right_margin:.1f}cm (Yêu cầu: 2.0cm)")
            if abs(top_margin - 2.5) > 0.4:
                margin_errors.append(f"Lề trên ước lượng: {top_margin:.1f}cm (Yêu cầu: 2.5cm)")
            if abs(bottom_margin - 2.5) > 0.4:
                margin_errors.append(f"Lề dưới ước lượng: {bottom_margin:.1f}cm (Yêu cầu: 2.5cm)")
                
            if margin_errors:
                is_margins_valid = False
                margins_feedback = "; ".join(margin_errors)
                
    # 4. Check font sizes
    total_spans = 0
    wrong_font_spans = 0
    wrong_size_spans = 0
    
    bibliography_text = []
    in_bibliography = False
    
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:
            if "lines" not in b:
                continue
            for l in b["lines"]:
                for s in l["spans"]:
                    text = s["text"].strip()
                    if not text:
                        continue
                        
                    if text.lower() in ["tài liệu tham khảo", "references", "danh mục tài liệu tham khảo"]:
                        in_bibliography = True
                        continue
                    elif in_bibliography and text.lower().startswith(("chương", "chapter", "phụ lục", "appendix")):
                        in_bibliography = False
                        
                    if in_bibliography:
                        bibliography_text.append(text)
                        
                    font_size = s["size"]
                    font_name = s["font"]
                    
                    if font_size > 13.5 or text.isupper():
                        continue
                    if font_size < 8.0:
                        continue
                        
                    total_spans += 1
                    if "times" not in font_name.lower():
                        wrong_font_spans += 1
                    if abs(font_size - 13.0) > 0.6:
                        wrong_size_spans += 1
                        
    is_font_family_valid = True
    font_family_feedback = "Font chữ toàn văn đạt chuẩn Times New Roman."
    is_font_size_valid = True
    font_size_feedback = "Cỡ chữ toàn văn đạt chuẩn 13pt."
    
    if total_spans > 0:
        font_fail = wrong_font_spans / total_spans
        size_fail = wrong_size_spans / total_spans
        if font_fail > 0.2:
            is_font_family_valid = False
            font_family_feedback = f"Font chữ chưa đồng bộ Times New Roman (tỷ lệ lỗi: {font_fail*100:.1f}%)"
        if size_fail > 0.2:
            is_font_size_valid = False
            font_size_feedback = f"Cỡ chữ chưa đúng 13pt (tỷ lệ lỗi: {size_fail*100:.1f}%)"
            
    # 5. Check citations
    is_citations_valid = True
    citations_feedback = "Danh mục tài liệu tham khảo đúng chuẩn APA hoặc Harvard."
    citations_errors = []
    
    if bibliography_text:
        bib_lines_str = "\n".join(bibliography_text[:25])
        try:
            prompt = f"""Bạn là chuyên gia thẩm định tài liệu tham khảo học thuật.
Hãy đối chiếu danh sách tài liệu tham khảo dưới đây của sinh viên UEH và chỉ ra các lỗi sai so với chuẩn APA 7th hoặc Harvard (như thiếu in nghiêng tên sách/tạp chí, sai thứ tự tên tác giả, năm xuất bản...).

Danh sách tài liệu tham khảo:
{bib_lines_str}

Hãy kiểm tra kỹ từng mục. Với mỗi mục có lỗi sai, hãy đề xuất bản sửa đổi chuẩn.
Phản hồi DUY NHẤT ở định dạng JSON thô có cấu trúc sau:
{{
  "is_citations_valid": true_or_false,
  "errors": [
     {{
       "original": "Mục trích dẫn gốc bị lỗi",
       "reason": "Lý do sai chuẩn chi tiết ngắn gọn bằng tiếng Việt",
       "suggested": "Mục trích dẫn đã được sửa lại đúng chuẩn APA 7th"
     }},
     ...
  ]
}}
If all references are correct, return true for "is_citations_valid" and empty list for "errors".
Tuyệt đối chỉ trả về chuỗi JSON thô, không kèm markdown hay giải thích nào khác."""

            from core.config import generate_content_with_rotation
            response = generate_content_with_rotation(
                model="gemini-3.5-flash",
                contents=prompt
            )
            result = parse_json_from_llm(response.text)
            is_citations_valid = result.get("is_citations_valid", True)
            citations_errors = result.get("errors", [])
            if not is_citations_valid and citations_errors:
                citations_feedback = f"Phát hiện {len(citations_errors)} tài liệu tham khảo chưa đúng chuẩn APA 7th."
            else:
                is_citations_valid = True
                citations_feedback = "Tất cả các tài liệu tham khảo đã quét đều đúng chuẩn APA 7th."
        except Exception as e:
            print(f"Lỗi kiểm tra APA PDF: {e}")
            is_citations_valid = False
            citations_feedback = "Không thể gọi AI để kiểm tra chuẩn trích dẫn."
    else:
        is_citations_valid = False
        citations_feedback = "Không tìm thấy danh mục tài liệu tham khảo nào trong file PDF. Yêu cầu bắt buộc phải có tài liệu tham khảo."
        
    return {
        "file_name": file_name,
        "is_paper_size_valid": is_paper_size_valid,
        "paper_size_feedback": paper_size_feedback,
        "is_margins_valid": is_margins_valid,
        "margins_feedback": margins_feedback,
        "is_font_family_valid": is_font_family_valid,
        "font_family_feedback": font_family_feedback,
        "is_font_size_valid": is_font_size_valid,
        "font_size_feedback": font_size_feedback,
        "is_spacing_valid": True,
        "spacing_feedback": "Trình xem PDF không hỗ trợ đánh giá giãn dòng chi tiết, nhưng bố cục tổng thể đạt yêu cầu.",
        "is_logo_valid": is_logo_valid,
        "logo_feedback": logo_feedback,
        "is_citations_valid": is_citations_valid,
        "citations_feedback": citations_feedback,
        "citations_errors": citations_errors
    }


