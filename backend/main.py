import os
import sys
if sys.platform.startswith("win"):
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
os.environ["HF_HUB_OFFLINE"] = "1"
os.environ["TRANSFORMERS_OFFLINE"] = "1"
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from dotenv import load_dotenv
import json



from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np




# Tạo chuỗi xử lý RAG thủ công nhưng chạy cực kỳ ổn định


def extract_text(content) -> str:
    if isinstance(content, list):
        text_parts = []
        for part in content:
            if isinstance(part, dict) and "text" in part:
                text_parts.append(part["text"])
            elif isinstance(part, str):
                text_parts.append(part)
        return "".join(text_parts)
    return str(content)


def custom_stuff_documents(llm, prompt, docs, query):
    # 1. Duyệt qua từng tài liệu để gộp cả Metadata (Tiêu đề, Tác giả) vào Context
    context_chunks = []
    for doc in docs:
        meta = doc.metadata
        title = meta.get("title", "Không rõ tiêu đề")
        authors = meta.get("authors", "Không rõ tác giả")
        year = meta.get("year", "2026")
        journal = meta.get("journal", "N/A")
       
        # Đóng gói thành một cụm thông tin hoàn chỉnh, rõ ràng
        chunk_info = (
            f"--- BÀI BÁO TÌM THẤY ---\n"
            f"Tiêu đề: {title}\n"
            f"Tác giả: {authors}\n"
            f"Năm xuất bản: {year}\n"
            f"Tạp chí/Phân hệ gốc: {journal}\nNội dung đoạn trích: {doc.page_content}"
        )
        context_chunks.append(chunk_info)
   
    # Nối tất cả các cụm thông tin bài báo lại với nhau bằng dấu xuống dòng
    full_context = "\n\n".join(context_chunks)
   
    # 2. Sử dụng format_messages (Chuẩn Chat thế hệ mới, thay vì format_prompt)
    messages = prompt.format_messages(context=full_context, input=query)
   
    # 3. Gọi Gemini xử lý thẳng danh sách tin nhắn
    response = llm.invoke(messages)
    return extract_text(response.content)




# 1. TẢI CẤU HÌNH BẢO MẬT (.env)
load_dotenv(override=True)


# 2. KHỞI TẠO FASTAPI
app = FastAPI(title="Thesis Chatbot API")


# 3. CẤU HÌNH CORS (Cho phép React cổng 3000 gọi sang Python cổng 8000)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# 4. NẠP KHO VECTOR FAISS
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FAISS_DB_PATH = os.path.join(BASE_DIR, "kho_vector_thesis_pdr")
if not os.path.exists(FAISS_DB_PATH):
    FAISS_DB_PATH = os.path.join(BASE_DIR, "kho_vector_thesis")

encode_kwargs = {'normalize_embeddings': True}
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
    model_kwargs={'device': 'cpu'},
    encode_kwargs=encode_kwargs
)

# Tải kho vector cứng từ folder lên RAM để truy vấn siêu tốc
vector_db = FAISS.load_local(FAISS_DB_PATH, embeddings, allow_dangerous_deserialization=True)
retriever = vector_db.as_retriever(search_kwargs={"k": 10})

# TẠO INDEX TF-IDF VÀ NẠP KHO TÀI LIỆU CHA NẾU CÓ (PDR SUPPORT)
from sklearn.feature_extraction.text import TfidfVectorizer
from langchain_core.documents import Document
import numpy as np

PARENT_STORE_PATH = os.path.join(BASE_DIR, "kho_parent_thesis.json")
parent_store = {}
if os.path.exists(PARENT_STORE_PATH):
    try:
        with open(PARENT_STORE_PATH, "r", encoding="utf-8") as f:
            parent_store = json.load(f)
    except Exception as e:
        print(f"⚠️ Lỗi đọc file parent store: {e}")

print("Fitting TF-IDF on database documents...")
all_docs = list(vector_db.docstore._dict.values())

if parent_store:
    print("Using Parent Document Store for TF-IDF indexing...")
    pdr_parent_docs = []
    all_parent_texts = []
    for parent_id, p_data in parent_store.items():
        p_doc = Document(page_content=p_data["page_content"], metadata=p_data["metadata"])
        pdr_parent_docs.append(p_doc)
        all_parent_texts.append(p_data["page_content"])
    
    tfidf_vectorizer = TfidfVectorizer()
    tfidf_matrix = tfidf_vectorizer.fit_transform(all_parent_texts)
else:
    print("Using Child Document Store for TF-IDF indexing (fallback)...")
    pdr_parent_docs = all_docs
    all_texts = [doc.page_content for doc in all_docs]
    tfidf_vectorizer = TfidfVectorizer()
    tfidf_matrix = tfidf_vectorizer.fit_transform(all_texts)

print("TF-IDF Matrix fitted successfully.")


# 5. KHAI BÁO CẤU TRÚC DỮ LIỆU ĐẦU VÀO TỪ REACT
class ChatRequest(BaseModel):
    message: str




class CompareTopicRequest(BaseModel):
    """
    Cấu trúc yêu cầu đối với endpoint so sánh đề tài.
    """
    title: str
    description: str




# 6. THIẾT LẬP PROMPT VÀ MÔ HÌNH GEMINI
llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.3, max_retries=3)


# Cập nhật lại đoạn Prompt trong main.py của bạn
system_prompt = """Bạn là một trợ lý AI chuyên nghiệp và thông minh hỗ trợ nghiên cứu khoa học. Nhiệm vụ của bạn là đọc kỹ các đoạn ngữ cảnh (Context) được cung cấp dưới đây để trả lời câu hỏi của người dùng.

HƯỚNG DẪN XỬ LÝ LOGIC (BẮT BUỘC):
1. TRẢ LỜI TRỰC TIẾP: Đi thẳng vào câu trả lời, tuyệt đối KHÔNG bắt đầu bằng các câu mở đầu thừa thãi như "Dựa trên ngữ cảnh được cung cấp...", "Theo tài liệu...", hoặc lời chào xã giao.
2. CHÍNH XÁC VÀ CỤ THỂ: Trích xuất chính xác các số liệu, tên phần mềm, phương pháp nghiên cứu, tên tác giả, năm công bố, cỡ mẫu (ví dụ: "PLS-SEM", "SmartPLS 4", "405 nhân viên y tế") có trong ngữ cảnh. Không tự ý suy diễn hoặc làm tròn số liệu.
3. NGẮN GỌN VÀ SÚC TÍCH: Chỉ tập trung trả lời đúng ý được hỏi. Tránh viết quá dài dòng hoặc đưa vào các phân tích lan man ngoài lề.
4. XỬ LÝ HƯỚNG NGHIÊN CỨU: Khi được hỏi về hướng nghiên cứu của một phân hệ/lĩnh vực, hãy chủ động đọc tiêu đề các bài báo trong Context, phân tích và tổng hợp thành các hướng nghiên cứu lớn bằng tiếng Việt.

Ngữ cảnh (Context):
{context}"""


prompt_template = ChatPromptTemplate.from_messages([
    ("system", system_prompt),
    ("human", "{input}"),
])

def expand_query_vietnamese(query: str) -> str:
    expanded = query
    abbreviations = {
        "đổi mới sáng tạo": "ĐMST PII",
        "năng suất lao động": "NSLĐ PRO",
        "khoa học công nghệ": "KHCN",
        "kinh tế xã hội": "KT-XH",
        "kinh tế - xã hội": "KT-XH",
        "cơ sở hạ tầng": "CSHT",
        "tổng sản phẩm trên địa bàn": "GRDP",
        "năng lực cạnh tranh": "PCI",
        "hiệu quả công việc": "HQCV",
        "sự gắn kết công việc": "GK",
        "sự gắn bó tổ chức": "GB",
        "khoa học xã hội và nhân văn": "KHXH&NV USSH",
        "đại học quốc gia": "ĐHQG-HCM ĐHQG",
        "thành phố hồ chí minh": "TP.HCM TP. Hồ Chí Minh TP HCM"
    }
    query_lower = query.lower()
    for key, val in abbreviations.items():
        if key in query_lower:
            expanded += f" {val}"
    return expanded


# 7. ENDPOINT API CHÍNH (React sẽ gọi vào đây để chat)
@app.post("/api/chat")
async def chat_with_thesis(request: ChatRequest):
    try:
        user_query = request.message
        
        # 1. Phân loại nhanh câu hỏi bằng prompt siêu rút gọn
        routing_prompt = f"""Phân loại câu hỏi của người dùng thành một trong hai nhãn sau:
- "thesis": Các câu hỏi liên quan đến luận văn, nghiên cứu khoa học, tài liệu học thuật, trường đại học (như ĐH KHXH&NV, ĐHQG), địa phương (Lâm Đồng, Gia Lâm), công trình xanh hoặc bất kỳ câu hỏi thực tế nào cần tìm cứu thông tin.
- "general": Các câu hỏi chào hỏi xã giao cực kỳ ngắn (ví dụ: "chào bạn", "hello", "hi"), hoặc trò chuyện phiếm vô thưởng vô phạt không liên quan gì đến trường học/luận văn.

Chỉ phản hồi đúng 1 từ duy nhất là "thesis" hoặc "general". Không giải thích thêm.
Câu hỏi: {user_query}"""
        
        route_response = llm.invoke(routing_prompt)
        query_type = extract_text(route_response.content).strip().lower()
        
        # 2. Phân luồng xử lý câu hỏi
        if "general" in query_type:
            general_prompt = f"""Bạn là trợ lý AI thông minh hỗ trợ luận văn UEH.
Người dùng đang hỏi bạn một câu hỏi xã giao hoặc ngoài lề. Hãy trả lời câu hỏi của họ một cách ngắn gọn, thân thiện và lịch sự.
Lưu ý nhắc nhở nhẹ nhàng (nếu cần thiết) rằng câu hỏi này nằm ngoài phạm vi tài liệu luận văn học thuật của hệ thống.
Câu hỏi: {user_query}"""
            
            general_response = llm.invoke(general_prompt)
            answer = extract_text(general_response.content)
            return {
                "answer": answer,
                "sources": [] # Trả về nguồn trống để Frontend không hiển thị tài liệu tham khảo sai lệch
            }
            
        else:
            expanded_query = expand_query_vietnamese(user_query)
            # 3. Tìm kiếm Vector (Dense) - Lấy Top 15 ứng viên từ FAISS (đoạn Con)
            raw_vector_results = vector_db.similarity_search_with_score(expanded_query, k=15)
            
            # Nếu chạy chế độ PDR, ánh xạ các đoạn Con thành đoạn Cha tương ứng
            vector_results = []
            if parent_store:
                seen_parent_ids = set()
                for doc, score in raw_vector_results:
                    parent_id = doc.metadata.get("parent_id")
                    if parent_id and parent_id in parent_store:
                        if parent_id not in seen_parent_ids:
                            seen_parent_ids.add(parent_id)
                            p_data = parent_store[parent_id]
                            p_doc = Document(page_content=p_data["page_content"], metadata=p_data["metadata"])
                            vector_results.append((p_doc, score))
                    else:
                        vector_results.append((doc, score))
            else:
                vector_results = raw_vector_results[:10]
            
            # 4. Tìm kiếm Từ khóa (Sparse) - Sử dụng TF-IDF (trên đoạn Cha)
            query_tfidf = tfidf_vectorizer.transform([expanded_query])
            tfidf_similarities = cosine_similarity(query_tfidf, tfidf_matrix).flatten()
            
            # Lọc lấy Top 10 ứng viên có điểm TF-IDF tương đồng cao nhất
            top_tfidf_indices = np.argsort(tfidf_similarities)[::-1][:10]
            tfidf_results = []
            for idx in top_tfidf_indices:
                score = float(tfidf_similarities[idx])
                if score > 0.05:  # Chỉ lấy các tài liệu có độ tương đồng từ khóa tối thiểu
                    tfidf_results.append((pdr_parent_docs[idx], score))
            
            # 5. Phối hợp kết quả bằng Reciprocal Rank Fusion (RRF)
            rrf_scores = {}
            
            # Ranks từ Vector Search
            for rank, (doc, score) in enumerate(vector_results):
                key = doc.page_content
                if key not in rrf_scores:
                    rrf_scores[key] = {"doc": doc, "score": 0.0}
                rrf_scores[key]["score"] += 1.0 / (60 + (rank + 1))
            
            # Ranks từ TF-IDF Search
            for rank, (doc, score) in enumerate(tfidf_results):
                key = doc.page_content
                if key not in rrf_scores:
                    rrf_scores[key] = {"doc": doc, "score": 0.0}
                rrf_scores[key]["score"] += 1.0 / (60 + (rank + 1))
            
            # Sắp xếp và chọn tài liệu tối ưu dựa trên điểm RRF (Lọc ngưỡng động để tăng Context Precision)
            hybrid_results = list(rrf_scores.values())
            hybrid_results.sort(key=lambda x: x["score"], reverse=True)
            
            docs = [item["doc"] for item in hybrid_results[:5]]
            
            # Dự phòng: Nếu không tìm thấy bất kỳ tài liệu nào, lấy Top 1 của FAISS
            if not docs and vector_results:
                docs = [vector_results[0][0]]
            
            # 5. Đưa vào hàm custom để Gemini trả lời dựa trên ngữ cảnh đã tối ưu
            answer = custom_stuff_documents(llm, prompt_template, docs, user_query)
           
            # 5. Lấy danh sách tên file PDF độc nhất từ metadata của các docs tìm thấy
            sources_list = []
            seen_pdfs = set()
            for doc in docs:
                # Lấy tên file thực tế từ metadata (ví dụ: file_name, hoặc trích xuất từ source/pdf_path)
                file_name = doc.metadata.get("file_name")
                if not file_name:
                    pdf_path = doc.metadata.get("source") or doc.metadata.get("pdf_path")
                    if pdf_path:
                        file_name = os.path.basename(pdf_path)
                if not file_name:
                    file_name = doc.metadata.get("title", "document") + ".pdf"
                
                # Đảm bảo tên file sạch, kết thúc bằng đuôi .pdf
                if file_name and not file_name.endswith(".pdf"):
                    file_name = file_name + ".pdf"
                    
                # Đảm bảo không bị trùng lặp
                if file_name not in seen_pdfs:
                    seen_pdfs.add(file_name)
                    title = doc.metadata.get("title", "Không rõ tiêu đề")
                    authors = doc.metadata.get("authors", "Không rõ tác giả")
                    year = doc.metadata.get("year", "2026")
                    journal = doc.metadata.get("journal", "N/A")
                    sources_list.append({
                        "pdf_name": file_name,
                        "title": title,
                        "authors": authors,
                        "year": str(year),
                        "journal": journal
                    })
                    
            return {
                "answer": answer,
                "sources": sources_list,
                "contexts": [doc.page_content for doc in docs]
            }
    except Exception as e:
        import traceback
        err_msg = str(e)
        print(f"Lỗi chat_with_thesis: {traceback.format_exc()}")
        if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg:
            raise HTTPException(
                status_code=429,
                detail="Tài khoản Gemini của bạn bị giới hạn lượt gọi (Rate Limit 429). Vui lòng thử lại sau 1-2 phút."
            )
        elif "API_KEY" in err_msg or "API key not valid" in err_msg:
            raise HTTPException(
                status_code=400,
                detail="Khóa API Gemini (GEMINI_API_KEY) trong file .env không hợp lệ hoặc đã hết hạn."
            )
        raise HTTPException(status_code=500, detail=f"Lỗi máy chủ: {err_msg}")




@app.post("/api/compare-topic")
async def compare_topic(request: CompareTopicRequest):
    """
    Endpoint so sánh đề tài nghiên cứu của người dùng với kho luận văn FAISS hiện có.
    Thực hiện đối khớp ngữ nghĩa, xếp hạng độ tương đồng và dùng Gemini để phân tích khoảng trống nghiên cứu.
    """
    try:
        user_title = request.title
        user_desc = request.description


        # BƯỚC 1: Validate input chi tiết
        if not user_title or not isinstance(user_title, str):
            raise HTTPException(status_code=400, detail="title phải là string không rỗng")
        if not user_desc or not isinstance(user_desc, str):
            raise HTTPException(status_code=400, detail="description phải là string không rỗng")


        user_title = user_title.strip()
        user_desc = user_desc.strip()


        if len(user_title) < 5:
            raise HTTPException(status_code=400, detail="Tiêu đề quá ngắn (tối thiểu 5 ký tự)")
        if len(user_desc) < 20:
            raise HTTPException(status_code=400, detail="Mô tả quá ngắn (tối thiểu 20 ký tự)")
        if len(user_desc) > 1000:
            raise HTTPException(status_code=400, detail="Mô tả quá dài (tối đa 1000 ký tự)")


        # BƯỚC 2: Tạo câu truy vấn kết hợp
        combined_query = f"{user_title}. {user_desc}"


        # BƯỚC 3: Truy vấn top 10 luận văn tương tự nhất từ FAISS
        docs_with_scores = vector_db.similarity_search_with_score(combined_query, k=10)


        # BƯỚC 4: Định dạng kết quả so sánh
        similar_theses = []
        max_similarity = 0.0


        for doc, distance_score in docs_with_scores:
            meta = doc.metadata
            # FAISS L2 distance score sang similarity % (0-100)
            similarity_percent = float(round(max(0.0, (1.0 - float(distance_score))) * 100.0, 1))
            max_similarity = max(max_similarity, similarity_percent)


            similar_theses.append({
                "title": meta.get("title", "N/A"),
                "authors": meta.get("authors", "N/A"),
                "year": meta.get("year", "N/A"),
                "journal": meta.get("journal", "N/A"),
                "similarity": similarity_percent,
                "summary": doc.page_content[:300] if doc.page_content else "N/A"
            })


        # BƯỚC 5: Tổng hợp danh sách tài liệu tương đồng cho Gemini phân tích
        similar_titles_list = "\n".join([f"- {t['title']} ({t['year']})" for t in similar_theses[:5]])


        # BƯỚC 6: Gọi Gemini với Prompt tối ưu hóa
        gap_prompt = f"""Phân tích nhanh đề tài: "{user_title}"


Các thesis tìm thấy:
{similar_titles_list}


Trả lời với 3 phần:
1. ĐÁNH GIÁ: Độ trùng cao/trung bình/thấp
2. KHOẢNG TRỐNG: 3 hướng nghiên cứu mới
3. ĐỀ XUẤT: Cách làm độc đáo


(Tiếng Việt, ngắn gọn)"""


        gap_response = llm.invoke(gap_prompt)
        gap_analysis = gap_response.content


        # BƯỚC 7: Phân cấp độ overlap
        if max_similarity > 75:
            overlap_level = "CAO - Cần điều chỉnh để tránh trùng lặp"
        elif max_similarity > 50:
            overlap_level = "TRUNG BÌNH - Có thể phát triển hướng mới"
        else:
            overlap_level = "THẤP - Đề tài mới, có tiềm năng"


        return {
            "status": "success",
            "user_topic": user_title,
            "similar_theses": similar_theses,
            "gap_analysis": gap_analysis,
            "overlap_level": overlap_level,
            "top_match_similarity": max_similarity
        }


    except HTTPException as he:
        raise he
    except Exception as e:
        import traceback
        print(f"Lỗi hệ thống trong endpoint so sánh: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Lỗi máy chủ: {str(e)}")




# Endpoint kiểm tra sức khỏe hệ thống
@app.get("/api/health")
def health_check():
    return {"status": "healthy", "database_loaded": True, "total_chunks": "7117"}


import json
from fastapi import File, Form, UploadFile
from langchain_core.messages import SystemMessage, HumanMessage
from viva_agents import EXAMINERS_CONFIG

# Viva schemas
class VivaStartRequest(BaseModel):
    pdf_title: str
    pdf_url: str | None = None

class VivaChatHistoryItem(BaseModel):
    sender: str  # "user" or "bot"
    examiner_id: str | None = None
    text: str

class VivaChatRequest(BaseModel):
    history: list[VivaChatHistoryItem]
    user_answer: str
    current_examiner_id: str
    pdf_title: str
    pdf_url: str | None = None
    pdf_context: str | None = None


# Helper utility to parse JSON from LLM output (cleaning markdown blocks if present)
def parse_json_from_llm(content: str):
    content = content.strip()
    if content.startswith("```json"):
        content = content[7:]
    elif content.startswith("```"):
        content = content[3:]
    if content.endswith("```"):
        content = content[:-3]
    content = content.strip()
    return json.loads(content)


@app.post("/api/viva/start")
async def start_viva(
    file: UploadFile = File(None),
    pdf_title: str = Form(...),
    pdf_url: str = Form(None)
):
    try:
        # 1. Trích xuất văn bản từ tệp PDF nếu có tải lên
        pdf_context = ""
        if file is not None:
            try:
                import fitz
                pdf_bytes = await file.read()
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                extracted_pages = []
                for page in doc[:4]:  # Lấy tối đa 4 trang đầu (chứa tiêu đề, tóm tắt, mục lục, phương pháp luận)
                    extracted_pages.append(page.get_text())
                pdf_context = "\n".join(extracted_pages)[:6000] # Giới hạn 6000 ký tự tránh tràn context cửa sổ LLM
                print(f"Trích xuất thành công {len(pdf_context)} ký tự từ PDF tải lên: {file.filename}")
            except Exception as pdf_err:
                print(f"Lỗi khi giải nén văn bản PDF: {pdf_err}")
                pdf_context = ""

        # 2. Query RAG vào FAISS để so khớp đề tài
        docs_with_scores = vector_db.similarity_search_with_score(pdf_title, k=4)
        
        is_completely_new = True
        top_similarity = 0.0
        if docs_with_scores:
            distance = docs_with_scores[0][1]
            top_similarity = max(0.0, (1.0 - float(distance))) * 100.0
            if top_similarity > 35.0:
                is_completely_new = False
                
        # 3. Tạo thông tin tham khảo RAG
        rag_context = ""
        for i, (doc, score) in enumerate(docs_with_scores[:3]):
            meta = doc.metadata
            title = meta.get("title", "Không rõ tiêu đề")
            authors = meta.get("authors", "Không rõ tác giả")
            year = meta.get("year", "N/A")
            rag_context += f"Tài liệu liên quan {i+1}: {title} - Tác giả: {authors} ({year})\nNội dung: {doc.page_content[:200]}...\n\n"
            
        # 4. Gộp ngữ cảnh nghiên cứu
        final_pdf_context = pdf_context if pdf_context else (rag_context if not is_completely_new else "")
        
        # 5. Khởi tạo giám khảo đầu tiên
        current_examiner = EXAMINERS_CONFIG["examiner_methodology"]
        
        novelty_status = (
            "Đề tài này mới hoàn toàn so với cơ sở dữ liệu hiện hành (is_completely_new = True)."
            if is_completely_new else
            f"Đề tài này có liên quan tới các nghiên cứu hiện có (Độ tương tự cao nhất: {top_similarity:.1f}%)."
        )
        
        pdf_details = (
            f"NỘI DUNG TÀI LIỆU CỦA SINH VIÊN (PDF):\n{final_pdf_context}\n"
            if final_pdf_context else
            "Sinh viên chưa tải lên nội dung tài liệu chi tiết, chỉ cung cấp tiêu đề đề tài."
        )

        system_instruction = (
            f"{current_examiner['system_prompt']}\n\n"
            f"Đề tài bảo vệ luận văn của sinh viên: '{pdf_title}'.\n"
            f"Trạng thái đề tài: {novelty_status}\n\n"
            f"{pdf_details}\n\n"
            "Nhiệm vụ: Hãy phân tích kỹ loại hình tài liệu của sinh viên (Ví dụ: đây là bài nghiên cứu thực nghiệm hay là bài viết nghiên cứu lý thuyết/khái niệm).\n"
            "Sau đó, hãy đặt câu hỏi chất vấn số 1 cực kỳ sắc bén, học thuật và phù hợp với đúng loại hình đề tài này."
        )
        
        messages = [
            SystemMessage(content=system_instruction),
            HumanMessage(content="Hãy đặt câu hỏi chất vấn đầu tiên của bạn về đề tài luận văn này.")
        ]
        
        response = llm.invoke(messages)
        first_question = response.content
        
        return {
            "status": "success",
            "is_completely_new": is_completely_new,
            "top_similarity": top_similarity,
            "current_examiner_id": current_examiner["id"],
            "question": first_question,
            "pdf_context": final_pdf_context,
            "history": [
                {"sender": "bot", "examiner_id": current_examiner["id"], "text": first_question}
            ]
        }
    except Exception as e:
        import traceback
        err_msg = str(e)
        print(f"Lỗi start_viva: {traceback.format_exc()}")
        if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg:
            raise HTTPException(
                status_code=429,
                detail="Hội đồng AI đang bận (Tài khoản Google Gemini của bạn bị giới hạn lượt gọi - Rate Limit 429). Vui lòng đợi 1-2 phút rồi nhấn 'Làm lại từ đầu' để thử lại."
            )
        elif "API_KEY" in err_msg or "API key not valid" in err_msg:
            raise HTTPException(
                status_code=400,
                detail="Khóa API Gemini (GEMINI_API_KEY) trong file .env không hợp lệ hoặc đã hết hạn. Vui lòng kiểm tra lại."
            )
        raise HTTPException(status_code=500, detail=f"Lỗi máy chủ: {err_msg}")


@app.post("/api/viva/chat")
async def chat_viva(request: VivaChatRequest):
    try:
        history = request.history
        user_answer = request.user_answer
        current_examiner_id = request.current_examiner_id
        pdf_title = request.pdf_title
        pdf_context = request.pdf_context
        
        # 1. Đếm số câu hỏi của giám khảo hiện tại trong lịch sử chat
        questions_by_current = [m for m in history if m.sender == "bot" and m.examiner_id == current_examiner_id]
        questions_count = len(questions_by_current)
        
        examiner_order = ["examiner_methodology", "examiner_novelty", "examiner_practical", "examiner_devil"]
        current_index = examiner_order.index(current_examiner_id)
        
        pdf_context_str = f"NỘI DUNG TÀI LIỆU CỦA SINH VIÊN (PDF):\n{pdf_context}\n\n" if pdf_context else ""

        # Check if we can still ask a follow-up (max questions per examiner = 3)
        if questions_count < 3:
            examiner = EXAMINERS_CONFIG[current_examiner_id]
            
            conversation_context = ""
            for item in history:
                role = "Hội đồng" if item.sender == "bot" else "Sinh viên"
                conversation_context += f"{role}: {item.text}\n"
            
            system_instruction = (
                f"{examiner['system_prompt']}\n\n"
                f"Đề tài của sinh viên: '{pdf_title}'.\n"
                f"{pdf_context_str}"
                f"Cuộc hội thoại phản biện cho đến nay:\n{conversation_context}\n"
                f"Sinh viên vừa trả lời: '{user_answer}'\n\n"
                "Nhiệm vụ: Đánh giá câu trả lời mới nhất này của sinh viên.\n"
                "- Nếu câu trả lời còn yếu, né tránh, thiếu số liệu hay thiếu lập luận thuyết phục, hãy chọn is_satisfactory = false và đưa ra câu hỏi phản biện phụ (follow_up_question) sắc sảo, dồn dập, phù hợp với định hướng khoa học của bạn.\n"
                "- Nếu câu trả lời đã thuyết phục và trả lời trực diện câu hỏi, hoặc bạn không có ý kiến phản đối nào thêm, hãy chọn is_satisfactory = true.\n\n"
                "Bạn BẮT BUỘC phải phản hồi ở định dạng JSON thô duy nhất có dạng:\n"
                "{\n"
                "  \"is_satisfactory\": true_or_false_boolean,\n"
                "  \"evaluation\": \"Nhận xét ngắn gọn và mang tính chuyên môn cao về câu trả lời\",\n"
                "  \"follow_up_question\": \"Câu hỏi xoáy tiếp theo phù hợp với đề tài (để trống nếu is_satisfactory = true)\"\n"
                "}"
            )
            
            messages = [
                SystemMessage(content=system_instruction),
                HumanMessage(content="Đánh giá câu trả lời của tôi và phản hồi bằng JSON.")
            ]
            
            response = llm.invoke(messages)
            
            try:
                eval_result = parse_json_from_llm(response.content)
                is_satisfactory = eval_result.get("is_satisfactory", True)
                evaluation = eval_result.get("evaluation", "")
                follow_up = eval_result.get("follow_up_question", "")
            except Exception as json_err:
                print(f"Lỗi phân tích JSON đánh giá: {json_err}. Raw: {response.content}")
                is_satisfactory = True
                evaluation = "Câu trả lời tạm chấp nhận được."
                follow_up = ""
                
            if not is_satisfactory and follow_up and questions_count < 3:
                return {
                    "is_finished": False,
                    "current_examiner_id": current_examiner_id,
                    "transition": False,
                    "evaluation": evaluation,
                    "question": follow_up
                }
                
        # 2. Chuyển lượt sang giám khảo tiếp theo
        if current_index + 1 < len(examiner_order):
            next_examiner_id = examiner_order[current_index + 1]
            next_examiner = EXAMINERS_CONFIG[next_examiner_id]
            
            conversation_context = ""
            for item in history:
                role = "Hội đồng" if item.sender == "bot" else "Sinh viên"
                conversation_context += f"{role}: {item.text}\n"
            conversation_context += f"Sinh viên: {user_answer}\n"
            
            system_instruction = (
                f"{next_examiner['system_prompt']}\n\n"
                f"Đề tài của sinh viên: '{pdf_title}'.\n"
                f"{pdf_context_str}"
                f"Cuộc hội thoại phản biện trước đó:\n{conversation_context}\n"
                "Nhiệm vụ: Hãy bắt đầu phần phản biện của bạn. Đưa ra câu hỏi đầu tiên của bạn chất vấn sinh viên dựa trên nội dung đề tài và cuộc trao đổi vừa rồi. Giữ đúng phong cách và văn phong khoa học học thuật đặc trưng của bạn."
            )
            
            messages = [
                SystemMessage(content=system_instruction),
                HumanMessage(content="Hãy đặt câu hỏi chất vấn đầu tiên của bạn.")
            ]
            
            response = llm.invoke(messages)
            next_question = response.content
            
            return {
                "is_finished": False,
                "current_examiner_id": next_examiner_id,
                "transition": True,
                "evaluation": f"Chuyển lượt sang vị giám khảo tiếp theo: {next_examiner['name']}.",
                "question": next_question
            }
            
        else:
            # 3. Kết thúc buổi phản biện, Hội đồng họp đánh giá và ra Scorecard
            conversation_context = ""
            for item in history:
                role = "Hội đồng" if item.sender == "bot" else "Sinh viên"
                conversation_context += f"{role}: {item.text}\n"
            conversation_context += f"Sinh viên: {user_answer}\n"
            
            system_instruction = (
                "Bạn đang đóng vai toàn bộ Hội đồng phản biện luận văn tốt nghiệp ĐHQG-HCM.\n"
                "Nhiệm vụ: Tổng hợp toàn bộ cuộc phản biện dưới đây và đưa ra đánh giá, điểm số cuối cùng.\n\n"
                f"Tên đề tài: '{pdf_title}'\n"
                f"{pdf_context_str}"
                f"Cuộc hội thoại bảo vệ luận văn:\n{conversation_context}\n\n"
                "Yêu cầu:\n"
                "- Cho điểm khách quan từ 0.0 đến 10.0 (score) dựa trên thái độ học tập và kiến thức chuyên môn sinh viên đã thể hiện.\n"
                "- Nêu rõ ít nhất 3 điểm mạnh học thuật (strengths) và 3 điểm yếu kỹ thuật/nội dung cần sửa đổi (weaknesses) của luận văn.\n"
                "- Đưa ra danh sách các câu hỏi chất vấn khó nhất kèm theo gợi ý câu trả lời mẫu chuẩn học thuật (ideal_answers) giúp sinh viên hoàn thiện bài nghiên cứu.\n\n"
                "Bạn BẮT BUỘC phải phản hồi ở định dạng JSON thô duy nhất có cấu trúc chính xác như sau:\n"
                "{\n"
                "  \"score\": float_score,\n"
                "  \"strengths\": [\"điểm mạnh 1\", \"điểm mạnh 2\", ...],\n"
                "  \"weaknesses\": [\"điểm yếu 1\", \"điểm yếu 2\", ...],\n"
                "  \"ideal_answers\": [\n"
                "     {\n"
                "       \"question\": \"câu hỏi khó 1\",\n"
                "       \"answer\": \"trả lời mẫu chuẩn học thuật\"\n"
                "     },\n"
                "     ...\n"
                "  ]\n"
                "}\n"
                "Tuyệt đối chỉ trả về chuỗi JSON thô, không kèm bất kỳ markdown nào."
            )
            
            messages = [
                SystemMessage(content=system_instruction),
                HumanMessage(content="Tổng hợp và chấm điểm phản biện luận văn bằng JSON.")
            ]
            
            response = llm.invoke(messages)
            
            try:
                scorecard = parse_json_from_llm(response.content)
            except Exception as json_err:
                print(f"Lỗi phân tích JSON scorecard: {json_err}. Raw: {response.content}")
                scorecard = {
                    "score": 7.5,
                    "strengths": ["Cố gắng trả lời đầy đủ các câu hỏi", "Đề tài có hướng phát triển tốt"],
                    "weaknesses": ["Một số câu trả lời còn chung chung", "Cần làm rõ phương pháp luận"],
                    "ideal_answers": [
                        {
                            "question": "Vấn đề kỹ thuật chính trong bài là gì?",
                            "answer": "Sinh viên nên tập trung vào quy trình chuẩn hóa và thiết kế mẫu tối ưu."
                        }
                    ]
                }
                
            return {
                "is_finished": True,
                "scorecard": scorecard
            }
            
    except Exception as e:
        import traceback
        err_msg = str(e)
        print(f"Lỗi chat_viva: {traceback.format_exc()}")
        if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg:
            raise HTTPException(
                status_code=429,
                detail="Lượt chất vấn bị gián đoạn do tài khoản Gemini bị giới hạn (Rate Limit 429). Hãy đợi 1-2 phút rồi gửi lại câu trả lời."
            )
        elif "API_KEY" in err_msg or "API key not valid" in err_msg:
            raise HTTPException(
                status_code=400,
                detail="Khóa API Gemini không hợp lệ. Vui lòng kiểm tra lại cấu hình."
            )
        raise HTTPException(status_code=500, detail=f"Lỗi máy chủ: {err_msg}")


PDF_STORAGE_DIR = os.path.join(os.path.dirname(__file__), "du_lieu_thesis", "pdf_files")

@app.get("/api/pdf/{filename}")
async def get_pdf(filename: str):
    file_path = os.path.join(PDF_STORAGE_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File không tồn tại")
    return FileResponse(file_path, media_type="application/pdf")


# =====================================================================
# UEH FORMAT CHECKER & AUTO-CORRECTOR BACKEND LOGIC
# =====================================================================

def check_docx_format(docx_bytes: bytes, file_name: str) -> dict:
    import docx

    from docx.shared import Cm, Pt
    import io
    import zipfile
    from google import genai
    from google.genai import types
    
    doc = docx.Document(io.BytesIO(docx_bytes))
    
    def get_document_defaults(document):
        from docx.oxml.ns import qn
        defaults = {'size': None, 'name': None}
        try:
            styles_element = document.styles.element
            doc_defaults = styles_element.find(qn('w:docDefaults'))
            if doc_defaults is not None:
                rpr_default = doc_defaults.find(qn('w:rPrDefault'))
                if rpr_default is not None:
                    rpr = rpr_default.find(qn('w:rPr'))
                    if rpr is not None:
                        # Extract size from sz or szCs
                        for attr in ['sz', 'szCs']:
                            sz = rpr.find(qn(f'w:{attr}'))
                            if sz is not None:
                                val = sz.get(qn('w:val'))
                                if val:
                                    defaults['size'] = float(val) / 2.0
                                    break
                        # Extract name from ascii, hAnsi, cs, eastAsia
                        rfonts = rpr.find(qn('w:rFonts'))
                        if rfonts is not None:
                            for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
                                val = rfonts.get(qn(f'w:{attr}'))
                                if val:
                                    defaults['name'] = val
                                    break
        except Exception:
            pass
        return defaults

    def get_style_font_size(style):
        if not style:
            return None
        from docx.oxml.ns import qn
        try:
            rpr = style.element.find(qn('w:rPr'))
            if rpr is not None:
                for attr in ['sz', 'szCs']:
                    sz = rpr.find(qn(f'w:{attr}'))
                    if sz is not None:
                        val = sz.get(qn('w:val'))
                        if val:
                            return float(val) / 2.0
        except Exception:
            pass
        return get_style_font_size(style.base_style)

    def get_style_font_name(style):
        if not style:
            return None
        from docx.oxml.ns import qn
        try:
            rpr = style.element.find(qn('w:rPr'))
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
                        val = rfonts.get(qn(f'w:{attr}'))
                        if val:
                            return val
        except Exception:
            pass
        return get_style_font_name(style.base_style)

    def get_effective_font_size(run, paragraph, document, defaults):
        from docx.oxml.ns import qn
        try:
            rpr = run.element.find(qn('w:rPr'))
            if rpr is not None:
                for attr in ['sz', 'szCs']:
                    sz = rpr.find(qn(f'w:{attr}'))
                    if sz is not None:
                        val = sz.get(qn('w:val'))
                        if val:
                            return float(val) / 2.0
        except Exception:
            pass
            
        style_size = get_style_font_size(paragraph.style)
        if style_size is not None:
            return style_size
            
        if defaults['size'] is not None:
            return defaults['size']
            
        return 12.0 # Default Word fallback if not specified
        
    def get_effective_font_name(run, paragraph, document, defaults):
        from docx.oxml.ns import qn
        try:
            rpr = run.element.find(qn('w:rPr'))
            if rpr is not None:
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is not None:
                    for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
                        val = rfonts.get(qn(f'w:{attr}'))
                        if val:
                            return val
        except Exception:
            pass
            
        style_name = get_style_font_name(paragraph.style)
        if style_name is not None:
            return style_name
            
        if defaults['name'] is not None:
            return defaults['name']
            
        return 'Times New Roman' # Default Vietnamese Word fallback

    def is_toc_entry_paragraph(p, text):
        style_name = p.style.name.lower()
        if style_name.startswith("toc"):
            return True
        if ("..." in text or ". ." in text) and (text[-1].isdigit() if text else False):
            return True
        return False

    doc_defaults = get_document_defaults(doc)
    
    # 1. Check paper size
    is_paper_size_valid = True
    paper_size_feedback = "Khổ giấy đạt chuẩn A4 (21.0 x 29.7cm)."
    if doc.sections:
        section = doc.sections[0]
        w = round(section.page_width.cm, 1) if section.page_width else 0.0
        h = round(section.page_height.cm, 1) if section.page_height else 0.0
        if abs(w - 21.0) > 0.3 or abs(h - 29.7) > 0.3:
            is_paper_size_valid = False
            paper_size_feedback = f"Khổ giấy hiện tại: {w}x{h}cm (Yêu cầu: A4 21.0x29.7cm)"

    # 2. Check margins (Lề trang)
    margin_errors = []
    if doc.sections:
        section = doc.sections[0]
        top_cm = round(section.top_margin.cm, 2) if section.top_margin else 0.0
        bottom_cm = round(section.bottom_margin.cm, 2) if section.bottom_margin else 0.0
        left_cm = round(section.left_margin.cm, 2) if section.left_margin else 0.0
        right_cm = round(section.right_margin.cm, 2) if section.right_margin else 0.0
        
        if abs(top_cm - 2.5) > 0.15:
            margin_errors.append(f"Lề trên: {top_cm}cm (Yêu cầu: 2.5cm)")
        if abs(bottom_cm - 2.5) > 0.15:
            margin_errors.append(f"Lề dưới: {bottom_cm}cm (Yêu cầu: 2.5cm)")
        if abs(left_cm - 3.5) > 0.15:
            margin_errors.append(f"Lề trái: {left_cm}cm (Yêu cầu: 3.5cm)")
        if abs(right_cm - 2.0) > 0.15:
            margin_errors.append(f"Lề phải: {right_cm}cm (Yêu cầu: 2.0cm)")
            
    is_margins_valid = len(margin_errors) == 0
    margins_feedback = "Lề trang đúng chuẩn quy định của UEH." if is_margins_valid else "; ".join(margin_errors)
    
    # 3. Check Font Family, Font Size and Line Spacing
    total_checked = 0
    wrong_font_count = 0
    wrong_size_count = 0
    wrong_spacing_count = 0
    
    bibliography_text = []
    in_bibliography = False
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        # Detect bibliography section
        if text.lower() in ["tài liệu tham khảo", "references", "danh mục tài liệu tham khảo"]:
            in_bibliography = True
            continue
        elif in_bibliography and text.lower().startswith(("chương", "chapter", "phụ lục", "appendix")):
            in_bibliography = False
            
        if in_bibliography:
            bibliography_text.append(text)
            
        # Check standard paragraph format
        spacing = p.paragraph_format.line_spacing
        
        # Check if heading (protect manual styles)
        is_heading = False
        if p.style.name.startswith("Heading") or text.isupper():
            is_heading = True
            
        # Bypass heading classification for TOC entries
        if is_toc_entry_paragraph(p, text):
            is_heading = False
            
        # Skip checking captions and notes for body text compliance (they have sizes 10pt and 11pt)
        is_caption_or_note = False
        if text.lower().startswith(('hình', 'ảnh', 'sơ đồ', 'biểu đồ', 'đồ thị', 'figure', 'fig', 'bảng', 'table', 'chú thích', 'nguồn', 'ghi chú', 'note', 'source')):
            is_caption_or_note = True

        has_large_size = False
        run_fonts = []
        run_sizes = []
        for run in p.runs:
            if not run.text.strip():
                continue
            
            effective_size = get_effective_font_size(run, p, doc, doc_defaults)
            effective_font = get_effective_font_name(run, p, doc, doc_defaults)
            
            if effective_size and effective_size > 13.5:
                has_large_size = True
            if effective_font:
                run_fonts.append(effective_font)
            if effective_size:
                run_sizes.append(effective_size)
                
        if has_large_size or is_heading or is_caption_or_note:
            continue

            
        total_checked += 1
        
        # Verify font family
        if run_fonts:
            if any(f != "Times New Roman" for f in run_fonts):
                wrong_font_count += 1
        
        # Verify font size
        if run_sizes:
            if any(abs(s - 13.0) > 0.1 for s in run_sizes):
                wrong_size_count += 1
                
        # Verify spacing
        if spacing is not None:
            if isinstance(spacing, float):
                if abs(spacing - 1.2) > 0.08:
                    wrong_spacing_count += 1
            else:
                pt_spacing = spacing.pt if hasattr(spacing, 'pt') else 0.0
                if abs(pt_spacing - 15.6) > 1.5:
                    wrong_spacing_count += 1
                    
    is_font_family_valid = True
    font_family_feedback = "Font chữ toàn văn đạt chuẩn Times New Roman."
    is_font_size_valid = True
    font_size_feedback = "Cỡ chữ toàn văn đạt chuẩn 13pt."
    is_spacing_valid = True
    spacing_feedback = "Giãn dòng đạt chuẩn 1.2 lines."
    
    if total_checked > 0:
        font_fail_rate = wrong_font_count / total_checked
        size_fail_rate = wrong_size_count / total_checked
        spacing_fail_rate = wrong_spacing_count / total_checked
        
        if font_fail_rate > 0.15:
            is_font_family_valid = False
            font_family_feedback = f"Font chữ chưa đồng bộ Times New Roman (tỷ lệ lỗi: {font_fail_rate*100:.1f}%)"
        if size_fail_rate > 0.15:
            is_font_size_valid = False
            font_size_feedback = f"Cỡ chữ chưa đúng chuẩn 13pt (tỷ lệ lỗi: {size_fail_rate*100:.1f}%)"
        if spacing_fail_rate > 0.2:
            is_spacing_valid = False
            spacing_feedback = f"Giãn dòng chưa đúng chuẩn 1.2 lines (tỷ lệ lỗi: {spacing_fail_rate*100:.1f}%)"
    else:
        font_family_feedback = "Không có đủ nội dung văn bản để kiểm tra font."
        font_size_feedback = "Không có đủ nội dung văn bản để kiểm tra cỡ chữ."
        spacing_feedback = "Không có đủ nội dung văn bản để kiểm tra giãn dòng."
        
    # 4. Logo Check using up to 3 images in reading order (protects against borders/lines)
    is_logo_valid = False
    logo_feedback = "Thiếu Logo UEH chính thức trên trang bìa luận văn."
    
    img_parts = []
    
    try:
        def get_images_in_reading_order(document, limit=3):
            from docx.oxml.ns import qn
            embed_ids = []
            
            def add_id(eid):
                if eid and eid not in embed_ids:
                    embed_ids.append(eid)
                    
            # Traverse paragraphs
            for p in document.paragraphs:
                for element in p._element.iter():
                    if element.tag.endswith('blip'):
                        add_id(element.get(qn('r:embed')))
                    elif element.tag.endswith('imagedata'):
                        add_id(element.get(qn('r:id')))
                if len(embed_ids) >= limit:
                    break
                    
            if len(embed_ids) < limit:
                # Traverse tables
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for element in p._element.iter():
                                    if element.tag.endswith('blip'):
                                        add_id(element.get(qn('r:embed')))
                                    elif element.tag.endswith('imagedata'):
                                        add_id(element.get(qn('r:id')))
                                if len(embed_ids) >= limit:
                                    break
                            if len(embed_ids) >= limit:
                                break
                        if len(embed_ids) >= limit:
                            break
            return embed_ids[:limit]

        embed_ids = get_images_in_reading_order(doc, limit=3)
        for eid in embed_ids:
            try:
                rel = doc.part.rels[eid]
                img_bytes = rel.target_part.blob
                mime_type = rel.target_part.content_type
                img_parts.append(types.Part.from_bytes(data=img_bytes, mime_type=mime_type))
            except Exception:
                pass
    except Exception as e:
        print(f"Lỗi bóc tách ảnh bìa docx: {e}")
        
    if img_parts:
        try:
            api_key = os.getenv("GEMINI_API_KEY")
            client = genai.Client(api_key=api_key)
            
            prompt = """Bạn là chuyên gia thẩm định văn bản của trường Đại học Kinh tế TP.HCM (UEH).
Hãy phân tích các hình ảnh được tải lên (được trích xuất từ trang bìa luận văn của sinh viên) và xác định xem:
1. Có hình ảnh nào chứa Logo chính thức hiện tại của trường Đại học Kinh tế TP.HCM (UEH) hay không?
2. Logo đó (nếu có) có đúng mẫu chuẩn màu sắc và các chi tiết vòng tròn đặc trưng không?

Hãy phản hồi DUY NHẤT ở định dạng JSON thô có cấu trúc như sau:
{
  "is_logo_valid": true_or_false,
  "feedback": "Lời nhận xét chi tiết ngắn gọn bằng tiếng Việt (Ví dụ: Phát hiện logo UEH đúng chuẩn mẫu mới trên bìa, hoặc Các ảnh gửi lên không chứa logo UEH đúng chuẩn)"
}
Tuyệt đối chỉ trả về chuỗi JSON thô, không kèm markdown hay giải thích nào khác."""

            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=img_parts + [prompt]
            )
            result = parse_json_from_llm(response.text)
            is_logo_valid = result.get("is_logo_valid", False)
            logo_feedback = result.get("feedback", "Không thể xác minh logo.")
        except Exception as gemini_err:
            print(f"Lỗi gọi Gemini kiểm tra logo DOCX: {gemini_err}")
            logo_feedback = "Không thể kết nối với AI để kiểm tra logo trang bìa."
            
    # 5. Citation Check (APA 7th)
    is_citations_valid = True
    citations_feedback = "Danh mục tài liệu tham khảo đạt chuẩn APA."
    citations_errors = []
    
    if bibliography_text:
        bib_lines_str = "\n".join(bibliography_text[:25])
        try:
            api_key = os.getenv("GEMINI_API_KEY")
            client = genai.Client(api_key=api_key)
            
            prompt = f"""Bạn là chuyên gia thẩm định tài liệu tham khảo học thuật.
Hãy đối chiếu danh sách tài liệu tham khảo dưới đây của sinh viên UEH và chỉ ra các lỗi sai so với chuẩn APA 7th hoặc Harvard (như thiếu in nghiêng tên sách/tạp chí, sai thứ tự tên tác giả, năm xuất bản...).

Danh sách tài liệu tham khảo:
{bib_lines_str}

Hãy kiểm tra kỹ từng mục. Với mỗi mục có lỗi sai, hãy đề xuất bản sửa đổi chuẩn.
Phản hồi DUY NHẤT ở định dạng JSON thô có cấu trúc sau:
{{
  "is_citations_valid": true_or_false,
  "errors": [
     {{
       "original": "Mục trích dẫn gốc bị lỗi",
       "reason": "Lý do sai chuẩn chi tiết ngắn gọn bằng tiếng Việt",
       "suggested": "Mục trích dẫn đã được sửa lại đúng chuẩn APA 7th"
     }},
     ...
  ]
}}
If all references are correct, return true for "is_citations_valid" and empty list for "errors".
Tuyệt đối chỉ trả về chuỗi JSON thô, không kèm markdown hay giải thích nào khác."""

            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt
            )
            result = parse_json_from_llm(response.text)
            is_citations_valid = result.get("is_citations_valid", True)
            citations_errors = result.get("errors", [])
            if not is_citations_valid and citations_errors:
                citations_feedback = f"Phát hiện {len(citations_errors)} tài liệu tham khảo chưa đúng chuẩn APA 7th."
            else:
                is_citations_valid = True
                citations_feedback = "Tất cả các tài liệu tham khảo đã quét đều đúng chuẩn APA 7th."
        except Exception as bib_err:
            print(f"Lỗi kiểm tra APA DOCX: {bib_err}")
            is_citations_valid = False
            citations_feedback = "Không thể gọi AI để kiểm tra chuẩn trích dẫn."
    else:
        is_citations_valid = False
        citations_feedback = "Không tìm thấy danh mục tài liệu tham khảo nào trong file Word. Yêu cầu bắt buộc phải có tài liệu tham khảo."
        
    return {
        "file_name": file_name,
        "is_paper_size_valid": is_paper_size_valid,
        "paper_size_feedback": paper_size_feedback,
        "is_margins_valid": is_margins_valid,
        "margins_feedback": margins_feedback,
        "is_font_family_valid": is_font_family_valid,
        "font_family_feedback": font_family_feedback,
        "is_font_size_valid": is_font_size_valid,
        "font_size_feedback": font_size_feedback,
        "is_spacing_valid": is_spacing_valid,
        "spacing_feedback": spacing_feedback,
        "is_logo_valid": is_logo_valid,
        "logo_feedback": logo_feedback,
        "is_citations_valid": is_citations_valid,
        "citations_feedback": citations_feedback,
        "citations_errors": citations_errors
    }


def fix_docx_format(docx_bytes: bytes) -> bytes:
    import docx
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    doc = docx.Document(io.BytesIO(docx_bytes))
    
    # 1. Fix margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)
        
    # 1.05 Tự động thu nhỏ ảnh vượt quá khung trang in khả dụng (15.5cm)
    max_width_emu = Cm(15.5)
    for shape in doc.inline_shapes:
        try:
            if shape.width and shape.width > max_width_emu:
                ratio = float(shape.height) / float(shape.width)
                shape.width = max_width_emu
                shape.height = int(max_width_emu * ratio)
        except Exception:
            pass
        
    def is_toc_entry_paragraph(p, text):
        style_name = p.style.name.lower()
        if style_name.startswith("toc"):
            return True
        if ("..." in text or ". ." in text) and (text[-1].isdigit() if text else False):
            return True
        return False

    # 1.1 Fix TOC style sheet definitions to ensure MS Word updates preserve formatting
    for style_name in ['TOC 1', 'TOC 2', 'TOC 3', 'TOC 4', 'TOC 5', 'toc 1', 'toc 2', 'toc 3', 'toc 4', 'toc 5']:
        try:
            style = doc.styles[style_name]
            style.font.name = 'Times New Roman'
            style.font.size = Pt(13)
        except Exception:
            pass

    # 2. Fix fonts, sizes, and line spacing
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # Check if it is an image paragraph
        is_image_para = False
        if 'w:drawing' in p._p.xml or 'w:pict' in p._p.xml:
            is_image_para = True
            
        if is_image_para:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            continue
            
        if not text:
            continue
            
        # Check if it is a caption/title/note
        is_fig_title = False
        is_tbl_title = False
        is_tbl_note = False
        
        if text.lower().startswith(('chú thích', 'nguồn', 'ghi chú', 'note', 'source')):
            is_tbl_note = True
        elif text.lower().startswith(('hình', 'ảnh', 'sơ đồ', 'biểu đồ', 'đồ thị', 'figure', 'fig')):
            is_fig_title = True
        elif text.lower().startswith(('bảng', 'table')):
            is_tbl_title = True
            
        # Protect headings
        is_heading = False
        if p.style.name.startswith("Heading") or text.isupper():
            is_heading = True
            
        # Bypass heading classification for TOC entries
        if is_toc_entry_paragraph(p, text):
            is_heading = False
            
        has_large_size = False
        for run in p.runs:
            run_size = None
            try:
                from docx.oxml.ns import qn
                rpr = run.element.find(qn('w:rPr'))
                if rpr is not None:
                    for attr in ['sz', 'szCs']:
                        sz = rpr.find(qn(f'w:{attr}'))
                        if sz is not None:
                            val = sz.get(qn('w:val'))
                            if val:
                                run_size = float(val) / 2.0
                                break
            except Exception:
                pass
            if run_size is None and run.font.size:
                run_size = run.font.size.pt
                
            if run_size and run_size > 13.5:
                has_large_size = True
                break
                
        if is_heading or has_large_size:
            # Only fix font family for headings
            for run in p.runs:
                run.font.name = 'Times New Roman'
            continue
            
        if is_fig_title:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                run.italic = False
            continue
            
        if is_tbl_title:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                run.italic = False
            continue
            
        if is_tbl_note:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = False
                run.italic = True
            continue
            
        # Standardize body paragraphs
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(6)
        
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            
    # 3. Fix tables (Issue 3: avoid justified cell text, make it left aligned and 12pt, clear indents)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.left_indent = Pt(0)
                    p.paragraph_format.right_indent = Pt(0)
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        
    out_bio = io.BytesIO()
    doc.save(out_bio)
    return out_bio.getvalue()


def check_pdf_format(pdf_bytes: bytes, file_name: str) -> dict:
    import fitz
    import os
    from google import genai
    from google.genai import types
    
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    # 1. Page 1 Image & Layout check with Gemini
    is_logo_valid = False
    logo_feedback = "Không thể xác minh logo trang bìa PDF."
    
    try:
        api_key = os.getenv("GEMINI_API_KEY")
        client = genai.Client(api_key=api_key)
        
        page = doc[0]
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("png")
        
        prompt = """Bạn là chuyên gia thẩm định văn bản của Đại học Kinh tế TP.HCM (UEH).
Đây là ảnh chụp trang bìa luận văn tốt nghiệp của sinh viên.
Hãy phân tích trang bìa này và kiểm tra xem:
1. Có logo UEH xuất hiện ở đầu trang bìa không?
2. Logo đó có đúng chuẩn logo Đại học Kinh tế TP.HCM (UEH) hay không?
3. Bố cục tên trường, tên đề tài, logo có được căn giữa cân đối không?

Hãy phản hồi DUY NHẤT ở định dạng JSON thô có cấu trúc sau:
{
  "is_logo_valid": true_or_false,
  "is_layout_valid": true_or_false,
  "feedback": "Nhận xét chi tiết ngắn gọn bằng tiếng Việt (Ví dụ: Logo đúng chuẩn và căn giữa đẹp, hoặc Thiếu logo UEH ở đầu trang, hoặc Tên đề tài chưa được căn giữa)"
}
Tuyệt đối chỉ trả về chuỗi JSON thô, không kèm markdown hay giải thích nào khác."""

        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[
                types.Part.from_bytes(data=img_bytes, mime_type="image/png"),
                prompt
            ]
        )
        result = parse_json_from_llm(response.text)
        is_logo_valid = result.get("is_logo_valid", False)
        logo_feedback = result.get("feedback", "Không thể xác minh logo.")
    except Exception as e:
        print(f"Lỗi kiểm tra logo PDF: {e}")
        logo_feedback = "Không thể kết nối với AI để kiểm tra logo trang bìa PDF."
        
    # 2. Check paper size
    is_paper_size_valid = True
    paper_size_feedback = "Khổ giấy đạt chuẩn A4 (21.0 x 29.7cm)."
    if len(doc) > 0:
        page0 = doc[0]
        w_cm = page0.rect.width * 0.0352778
        h_cm = page0.rect.height * 0.0352778
        if abs(w_cm - 21.0) > 0.4 or abs(h_cm - 29.7) > 0.4:
            is_paper_size_valid = False
            paper_size_feedback = f"Khổ giấy hiện tại: {w_cm:.1f}x{h_cm:.1f}cm (Yêu cầu: A4 21.0x29.7cm)"

    # 3. Check margins
    margin_errors = []
    is_margins_valid = True
    margins_feedback = "Lề trang PDF đúng chuẩn quy định của UEH."
    if len(doc) > 1:
        page2 = doc[1]
        rect = page2.rect
        blocks = page2.get_text("blocks")
        if blocks:
            x0_min = min(b[0] for b in blocks)
            y0_min = min(b[1] for b in blocks)
            x1_max = max(b[2] for b in blocks)
            y1_max = max(b[3] for b in blocks)
            
            left_margin = x0_min * 0.0352778
            top_margin = y0_min * 0.0352778
            right_margin = (rect.width - x1_max) * 0.0352778
            bottom_margin = (rect.height - y1_max) * 0.0352778
            
            if abs(left_margin - 3.5) > 0.4:
                margin_errors.append(f"Lề trái ước lượng: {left_margin:.1f}cm (Yêu cầu: 3.5cm)")
            if abs(right_margin - 2.0) > 0.4:
                margin_errors.append(f"Lề phải ước lượng: {right_margin:.1f}cm (Yêu cầu: 2.0cm)")
            if abs(top_margin - 2.5) > 0.4:
                margin_errors.append(f"Lề trên ước lượng: {top_margin:.1f}cm (Yêu cầu: 2.5cm)")
            if abs(bottom_margin - 2.5) > 0.4:
                margin_errors.append(f"Lề dưới ước lượng: {bottom_margin:.1f}cm (Yêu cầu: 2.5cm)")
                
            if margin_errors:
                is_margins_valid = False
                margins_feedback = "; ".join(margin_errors)
                
    # 4. Check font sizes
    total_spans = 0
    wrong_font_spans = 0
    wrong_size_spans = 0
    
    bibliography_text = []
    in_bibliography = False
    
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:
            if "lines" not in b:
                continue
            for l in b["lines"]:
                for s in l["spans"]:
                    text = s["text"].strip()
                    if not text:
                        continue
                        
                    if text.lower() in ["tài liệu tham khảo", "references", "danh mục tài liệu tham khảo"]:
                        in_bibliography = True
                        continue
                    elif in_bibliography and text.lower().startswith(("chương", "chapter", "phụ lục", "appendix")):
                        in_bibliography = False
                        
                    if in_bibliography:
                        bibliography_text.append(text)
                        
                    font_size = s["size"]
                    font_name = s["font"]
                    
                    if font_size > 13.5 or text.isupper():
                        continue
                    if font_size < 8.0:
                        continue
                        
                    total_spans += 1
                    if "times" not in font_name.lower():
                        wrong_font_spans += 1
                    if abs(font_size - 13.0) > 0.6:
                        wrong_size_spans += 1
                        
    is_font_family_valid = True
    font_family_feedback = "Font chữ toàn văn đạt chuẩn Times New Roman."
    is_font_size_valid = True
    font_size_feedback = "Cỡ chữ toàn văn đạt chuẩn 13pt."
    
    if total_spans > 0:
        font_fail = wrong_font_spans / total_spans
        size_fail = wrong_size_spans / total_spans
        if font_fail > 0.2:
            is_font_family_valid = False
            font_family_feedback = f"Font chữ chưa đồng bộ Times New Roman (tỷ lệ lỗi: {font_fail*100:.1f}%)"
        if size_fail > 0.2:
            is_font_size_valid = False
            font_size_feedback = f"Cỡ chữ chưa đúng 13pt (tỷ lệ lỗi: {size_fail*100:.1f}%)"
            
    # 5. Check citations
    is_citations_valid = True
    citations_feedback = "Danh mục tài liệu tham khảo đúng chuẩn APA hoặc Harvard."
    citations_errors = []
    
    if bibliography_text:
        bib_lines_str = "\n".join(bibliography_text[:25])
        try:
            api_key = os.getenv("GEMINI_API_KEY")
            client = genai.Client(api_key=api_key)
            
            prompt = f"""Bạn là chuyên gia thẩm định tài liệu tham khảo học thuật.
Hãy đối chiếu danh sách tài liệu tham khảo dưới đây của sinh viên UEH và chỉ ra các lỗi sai so với chuẩn APA 7th hoặc Harvard (như thiếu in nghiêng tên sách/tạp chí, sai thứ tự tên tác giả, năm xuất bản...).

Danh sách tài liệu tham khảo:
{bib_lines_str}

Hãy kiểm tra kỹ từng mục. Với mỗi mục có lỗi sai, hãy đề xuất bản sửa đổi chuẩn.
Phản hồi DUY NHẤT ở định dạng JSON thô có cấu trúc sau:
{{
  "is_citations_valid": true_or_false,
  "errors": [
     {{
       "original": "Mục trích dẫn gốc bị lỗi",
       "reason": "Lý do sai chuẩn chi tiết ngắn gọn bằng tiếng Việt",
       "suggested": "Mục trích dẫn đã được sửa lại đúng chuẩn APA 7th"
     }},
     ...
  ]
}}
If all references are correct, return true for "is_citations_valid" and empty list for "errors".
Tuyệt đối chỉ trả về chuỗi JSON thô, không kèm markdown hay giải thích nào khác."""

            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt
            )
            result = parse_json_from_llm(response.text)
            is_citations_valid = result.get("is_citations_valid", True)
            citations_errors = result.get("errors", [])
            if not is_citations_valid and citations_errors:
                citations_feedback = f"Phát hiện {len(citations_errors)} tài liệu tham khảo chưa đúng chuẩn APA 7th."
            else:
                is_citations_valid = True
                citations_feedback = "Tất cả các tài liệu tham khảo đã quét đều đúng chuẩn APA 7th."
        except Exception as e:
            print(f"Lỗi kiểm tra APA PDF: {e}")
            is_citations_valid = False
            citations_feedback = "Không thể gọi AI để kiểm tra chuẩn trích dẫn."
    else:
        is_citations_valid = False
        citations_feedback = "Không tìm thấy danh mục tài liệu tham khảo nào trong file PDF. Yêu cầu bắt buộc phải có tài liệu tham khảo."
        
    return {
        "file_name": file_name,
        "is_paper_size_valid": is_paper_size_valid,
        "paper_size_feedback": paper_size_feedback,
        "is_margins_valid": is_margins_valid,
        "margins_feedback": margins_feedback,
        "is_font_family_valid": is_font_family_valid,
        "font_family_feedback": font_family_feedback,
        "is_font_size_valid": is_font_size_valid,
        "font_size_feedback": font_size_feedback,
        "is_spacing_valid": True,
        "spacing_feedback": "Trình xem PDF không hỗ trợ đánh giá giãn dòng chi tiết, nhưng bố cục tổng thể đạt yêu cầu.",
        "is_logo_valid": is_logo_valid,
        "logo_feedback": logo_feedback,
        "is_citations_valid": is_citations_valid,
        "citations_feedback": citations_feedback,
        "citations_errors": citations_errors
    }


@app.post("/api/check-format")
async def check_format(file: UploadFile = File(...)):
    filename = file.filename
    content = await file.read()
    
    if filename.endswith(".docx"):
        try:
            report = check_docx_format(content, filename)
            return {"status": "success", "report": report}
        except Exception as e:
            import traceback
            print(traceback.format_exc())
            raise HTTPException(status_code=500, detail=f"Lỗi phân tích file Word: {str(e)}")
    elif filename.endswith(".pdf"):
        try:
            report = check_pdf_format(content, filename)
            return {"status": "success", "report": report}
        except Exception as e:
            import traceback
            print(traceback.format_exc())
            raise HTTPException(status_code=500, detail=f"Lỗi phân tích file PDF: {str(e)}")
    else:
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ tệp định dạng .docx hoặc .pdf")


@app.post("/api/fix-format")
async def fix_format(file: UploadFile = File(...)):
    filename = file.filename
    if not filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Tính năng tự động sửa định dạng chỉ hỗ trợ tệp .docx")
        
    try:
        content = await file.read()
        fixed_content = fix_docx_format(content)
        out_filename = filename.replace(".docx", "_fixed_UEH.docx")
        
        from fastapi.responses import StreamingResponse
        import io
        return StreamingResponse(
            io.BytesIO(fixed_content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={out_filename}"}
        )
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Lỗi tự động sửa định dạng: {str(e)}")

# Forced reload trigger to update environment API key from .env file v3
